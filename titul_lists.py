import settings
import docx;
import copy;
import pandas as pd;

#region titul_lists
def add_titul(main_document,path_to_titul,path_to_emblem,disc_name, autor="au"):
    titul = docx.Document(path_to_titul);
    titul.add_page_break()
    titul.tables[2].cell(0,0).text = "";
    #titul.tables[3].cell(0,1).text=""

    set_formatting(titul.tables[2].cell(0,0).paragraphs[0].add_run(disc_name))
    titul.tables[2].cell(0,0).paragraphs[0].alignment = 1
    #set_formatting(titul.tables[3].cell(0,1).paragraphs[0].add_run(autor))
    add_picture(main_document,path_to_emblem)
    for element in titul.element.body:
        main_document.element.body.append(copy.copy(element))

def set_formatting(temp):
        temp.font.bold=True;
        temp.font.name='Times New Roman';
        temp.font.size=docx.shared.Pt(14);


def add_picture(main_document,path_to_emblem):
    if (len(main_document.paragraphs)!=0):
        add_picture_to_end(main_document,path_to_emblem)
    else:
        main_document.add_paragraph();
        add_picture_to_end(main_document,path_to_emblem)

def add_picture_to_end(main_document,path_to_emblem):
        p=main_document.paragraphs[-1];
        p.alignment = 1
        r=p.add_run();
        r.add_picture(path_to_emblem)

def titul_lists():
    df=pd.read_excel(settings.path_to_plan, sheet_name='ПланСвод')
    disc_names=df.iloc[:,2]
    autor="ss"
    main_document = docx.Document();
    iter=0;
    counter=0
    while iter<len(disc_names):
        disc_name = disc_names[iter];
        if (type(disc_name)==str and disc_name!="Наименование" and df.iloc[iter,23]=='160'):
            counter+=1;
            print(counter)
            #add_titul(main_document,path_to_titul,path_to_emblem,df.iloc[iter,1]+" "+disc_name, autor)
            add_titul(main_document,settings.path_to_titul,settings.path_to_emblem,disc_name, autor)
        iter=iter+1
    main_document.save(settings.path_to_doc);
#endregion





if __name__ == "__main__":
    titul_lists();


  