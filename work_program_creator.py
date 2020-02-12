import docx
import pandas as pd
import settings
import re
import copy
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph 
import random

def parse_code(code):
    vari=False;
    dv=False
    import re;
    if (len(re.findall('.О.',code))!=0):
        return 'обязательной дисциплиной'
    if (len(re.findall('.В.',code))!=0):
        vari=True
        if (len(re.findall('.ДВ.',code))!=0):
            dv=True
    if (vari):
        if dv:
            return 'дисциплиной по выбору вариативной части'
        else:
            return 'дисциплиной вариативной части'
    else:
        return 'дисциплиной'
    
def format_row(row,font_size,ali=WD_TABLE_ALIGNMENT.LEFT):
    column_number=0
    for cell in row.cells:
        for paragrap in cell.paragraphs:
            paragrap.alignment=ali
            for run in paragrap.runs:
                run.font.size = docx.shared.Pt(font_size);
                run.font.name='Times New Roman';
        column_number+=1;
        
def format_cell(cell,font_size,ali=WD_TABLE_ALIGNMENT.CENTER):
    for paragrap in cell.paragraphs:
        paragrap.alignment=ali
        for run in paragrap.runs:
            run.font.size = docx.shared.Pt(font_size);
            run.font.name='Times New Roman';
    
def get_week(length,i):
    k=15/length
#for i in indexes:
    week1=int(i*k+1)
    week2=int((i+1)*k+1)
    if week1!=week2:
        week=str(week1)+'-'+str(week2)
    else:
        week=str(week1)
    return week

def add_text_content(doc,pattern,text):
    indexes = range(0,len(doc.paragraphs))
    for i in indexes:
        if pattern+'b' in doc.paragraphs[i].text:
            doc.paragraphs[i].text=doc.paragraphs[i].text.replace(pattern+'b',text)
            for r in doc.paragraphs[i].runs:
                r.bold=True
                r.font.name='Times New Roman';
                r.font.size=docx.shared.Pt(14);
        if pattern in doc.paragraphs[i].text:
            doc.paragraphs[i].text=doc.paragraphs[i].text.replace(pattern,text)
            for r in doc.paragraphs[i].runs:
                #r.bold=True
                r.font.name='Times New Roman';
                r.font.size=docx.shared.Pt(14);
                
def add_paragraphs_content(doc,pattern,text,row_start=None):
            
    indexes = range(0,len(doc.paragraphs))
    for i in indexes:
        if pattern in doc.paragraphs[i].text:
            add_counter=0
            for row in reversed(text):
                add_counter+=1    
                if (row_start is None):
                    doc.paragraphs[i].insert_paragraph_before(str(len(text)-add_counter+1)+'. '+row)
                else:
                    doc.paragraphs[i].insert_paragraph_before(str(row_start)+' '+row)
                for r in doc.paragraphs[i].runs:
                    r.font.name='Times New Roman';
                    r.font.size=docx.shared.Pt(14);
            doc.paragraphs[i+add_counter]._element.getparent().remove(doc.paragraphs[i+add_counter]._element)   
            add_paragraphs_content(doc,pattern,text,row_start)
            break;

def remove_paragraphs_by_pattern(doc,pattern):
    indexes = range(0,len(doc.paragraphs))
    for i in indexes:
        if pattern in doc.paragraphs[i].text:
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)   
            remove_paragraphs_by_pattern(doc,pattern)
            break;
    
def remove_pattern(doc,pattern):
    indexes = range(0,len(doc.paragraphs))
    for i in indexes:
        if pattern in doc.paragraphs[i].text:
            doc.paragraphs[i].text=doc.paragraphs[i].text.replace(pattern,'')
            for r in doc.paragraphs[i].runs:
                r.font.name='Times New Roman';
                r.font.size=docx.shared.Pt(14);
            remove_pattern(doc,pattern)
            break;

                
def set_na_for_table_6_4(temp_table,column_number,natext):
    temp_table.cell(2,column_number).text=natext;
    temp_table.cell(3,column_number).text=natext;
    temp_table.cell(4,column_number).text=natext;
    temp_table.cell(5,column_number).text=natext;
    temp_table.cell(6,column_number).text=natext;
    temp_table.cell(7,column_number).text=natext;
    
    
def add_content_to_tables(disc,doc):
    indexes = range(0,len(doc.tables))
    for_remove=[]
    for i in indexes:
        temp_table=doc.tables[i];

        if (temp_table.cell(0,0).text=='Процедура проведения'):
            attestation = disc.get_attestation();
            natext='-'
            if (not disc.lab_check()):
                set_na_for_table_6_4(temp_table,3,natext);
            if (not disc.pract_check()):
                set_na_for_table_6_4(temp_table,2,natext)
            if (not attestation.examen):
                set_na_for_table_6_4(temp_table,4,natext)
            if (not attestation.zachet_s_ocenk):
                set_na_for_table_6_4(temp_table,5,natext)
            if (not attestation.zachet):
                set_na_for_table_6_4(temp_table,6,natext)
            if (not attestation.course_work):
                set_na_for_table_6_4(temp_table,7,natext)
            for row in temp_table.rows:
                format_row(row,8,WD_TABLE_ALIGNMENT.CENTER)
            temp_table.style='Table Grid'
            
        if (temp_table.cell(0,0).text=='Формируемые компетенции (код и название компетенции, уровень освоения - при наличии в карте компетенции)'):
            comp_number=0
            for comp in disc.competentions:
                row = temp_table.add_row()
                row.cells[1].text = "Знать "+comp.znat
                format_row(row,14)
                row = temp_table.add_row()
                row.cells[1].text = "Уметь "+comp.umet
                format_row(row,14)
                row = temp_table.add_row()
                row.cells[1].text = "Владеть "+comp.vladet
                format_row(row,14)
                temp_table.cell(comp_number+1,0).merge(temp_table.cell(comp_number+2,0))
                temp_table.cell(comp_number+2,0).merge(temp_table.cell(comp_number+3,0))
                
                
                row.cells[0].text=comp.code+" ("+comp.name+")"
                format_row(row,14)
                comp_number+=3
                
                
            temp_table.style='Table Grid'
        
        if (temp_table.cell(0,0).text=='Элементы компетенций (знания, умения, владения) '):
            comp_number=0
            textcontrol="Текущий контроль: Выполнение устных заданий; "
            if disc.lab_check():
                textcontrol+="Выполнение лабораторных работ; "
            if disc.pract_check():
                textcontrol+="Выполнение практических заданий; "
            
            attestation = disc.get_attestation();
            if attestation.course_work:
                textcontrol+="Курсовая работа; "
            if attestation.examen:
                textcontrol+="Экзамен; "               
            if attestation.zachet:
                textcontrol+="Зачет; "    
            if attestation.zachet_s_ocenk:
                textcontrol+="Зачет с оценкой; "    
                               

            for comp in disc.competentions:
                
                row = temp_table.add_row()
                row.cells[0].text = "Знать ("+comp.code+")"
                row.cells[1].text = "Знать "+comp.znat
                row.cells[2].text = "Правильность и полнота ответов, глубина понимания вопроса"
                row.cells[3].text = textcontrol
                row.cells[4].text = "Шкала 1"
                format_row(row,12)
                row = temp_table.add_row()
                row.cells[0].text = "Уметь ("+comp.code+")"
                row.cells[1].text = "Уметь "+comp.umet
                row.cells[2].text = "Правильность выполнения учебных заданий, аргументированность выводов"
                row.cells[3].text = textcontrol
                row.cells[4].text = "Шкала 1"
                format_row(row,12)
                row = temp_table.add_row()
                row.cells[0].text = "Владеть ("+comp.code+")"
                row.cells[1].text = "Владеть "+comp.vladet
                row.cells[2].text = "Обоснованность и аргументированность выполнения учебной деятельности"
                row.cells[3].text = textcontrol
                row.cells[4].text = "Шкала 2"
                format_row(row,12)
                
                
            temp_table.style='Table Grid'
        
        
        if (temp_table.cell(0,0).text=='№ раздела' 
            and temp_table.cell(0,1).text=='Наименование раздела' 
            and temp_table.cell(0,2).text=='Содержание раздела'):
            count=0;
            for sem in disc.semesters:
                for thema in sem.themes:
                    row = temp_table.add_row()
                    count+=1
                    row.cells[0].text=str_(thema.number);
                    row.cells[1].text=thema.name;
                    row.cells[2].text=thema.contain;
                    format_row(row,14)
            temp_table.style='Table Grid'
            
        if (temp_table.cell(0,0).text=='№ п/п' 
            and temp_table.cell(0,1).text=='№ раздела дисциплины' 
            and temp_table.cell(0,2).text=='Наименование лабораторных работ'
            and temp_table.cell(0,3).text=='Трудоемкость (в акад. часах)'):
            if (disc.lab_check()):
                count=0;
                su0=0
                for sem in disc.semesters:
                    su=0
                    for thema in sem.themes:
                        for laba in thema.lab_works:
                            row = temp_table.add_row()
                            count+=1
                            row.cells[0].text=str_(count);
                            row.cells[1].text=str_(thema.number);
                            row.cells[2].text=laba.theme;
                            row.cells[3].text=str_(laba.hours);
                            format_row(row,14)
                            su+=laba.hours
                            su0+=laba.hours
                    if su!=0:
                        row = temp_table.add_row()
                        row.cells[2].text="Всего в "+str_(sem.number)+" семестре";
                        row.cells[3].text=str_(su);
                        format_row(row,14)
                if su0!=0:
                    row = temp_table.add_row()
                    row.cells[2].text="Всего";
                    row.cells[3].text=str_(su0);
                    format_row(row,14)
                temp_table.style='Table Grid'
            else:
                for_remove.append(temp_table)
                
            
        if (temp_table.cell(0,0).text=='№ п/п' 
            and temp_table.cell(0,1).text=='№ раздела дисциплины' 
            and temp_table.cell(0,2).text=='Тематика практических занятий'
            and temp_table.cell(0,3).text=='Трудоемкость (в акад. часах)'):
            if (disc.pract_check()):
                count=0;
                su0=0
                for sem in disc.semesters:
                    su=0
                    for thema in sem.themes:
                        for pract in thema.practicals:
                            if (pract.hours!=0):
                                row = temp_table.add_row()
                                count+=1
                                row.cells[0].text=str_(count);
                                row.cells[1].text=str_(thema.number);
                                row.cells[2].text=pract.contain;
                                su+=pract.hours
                                su0+=pract.hours
                                row.cells[3].text=str_(pract.hours);
                                format_row(row,14)
                    if su!=0:
                        row = temp_table.add_row()
                        row.cells[2].text="Всего в "+str_(sem.number)+" семестре";
                        row.cells[3].text=str_(su);
                        format_row(row,14)
                if su0!=0:
                    row = temp_table.add_row()
                    row.cells[2].text="Всего";
                    row.cells[3].text=str_(su0);
                    format_row(row,14)
                temp_table.style='Table Grid'
            else:
                for_remove.append(temp_table)

        if (temp_table.cell(0,0).text=='№ раздела' 
            and temp_table.cell(0,1).text=='Семестр' 
            and temp_table.cell(0,2).text=='Неделя семестра'):

            le_ss=0
            la_ss=0
            pa_ss=0
            sa_ss=0
            co_ss=0
            kr_pa_ss=0;
            for sem in disc.semesters:
                control=''
                if sem.attestation.examen:
                    control+='экзамен'
                if sem.attestation.zachet:
                    if (len(control)>1):
                        control+=', зачет'
                    else:
                        control+='зачет'
                if sem.attestation.zachet_s_ocenk:
                    if (len(control)>1):
                        control+=', зачет с оценкой'
                    else:
                        control+='зачет с оценкой'
                if sem.attestation.course_work:
                    if (len(control)>1):
                        control+=', курс. раб'
                    else:
                        control+='курс. раб'
                it=0;
                le_s=0
                la_s=0
                pa_s=0
                sa_s=0
                co_s=0
                kr_pa_s=0;
                for theme in sem.themes:
                    temp_con='Устное собеседование'
                    row = temp_table.add_row()
                    le = sem.hours_summary.lections/len(sem.themes)
                    la = theme.get_lab_hours()
                    pa = theme.get_pract_hours()
                    sa= sem.hours_summary.self_work/len(sem.themes)
                    le_s+=le
                    la_s+=la
                    pa_s+=pa
                    sa_s+=sa

                    if la>0:
                        temp_con+=', защита лаб. работы'
                    if pa>0:
                        temp_con+=', выполнение практ. заданий'
                    row.cells[5].text=prepair_value_to_str(le)
                    row.cells[6].text=prepair_value_to_str(la)
                    row.cells[7].text=prepair_value_to_str(pa)
                    row.cells[4].text=prepair_value_to_str(le+la+pa)
                    row.cells[3].text=prepair_value_to_str(le+la+pa+sa)
                    row.cells[8].text=prepair_value_to_str(sa)
                    row.cells[1].text=str_(sem.number)
                    row.cells[2].text=get_week(len(sem.themes),it);
                    it+=1
                    row.cells[0].text=str_(theme.number)
                    row.cells[11].text=temp_con;
                    format_row(row,9)
                row = temp_table.add_row()
                row.cells[0].merge(row.cells[1])
                row.cells[1].merge(row.cells[2])
                row.cells[2].text="По материалам "+str_(sem.number)+" семестра";
                co_s+=sem.hours_summary.control;
                row.cells[10].text=prepair_value_to_str(sem.hours_summary.control);
                row.cells[11].text=control;
                row.cells[9].text=prepair_value_to_str(sem.krpa);
                format_row(row,9)
                kr_pa_s+=sem.krpa;
                row = temp_table.add_row()
                row.cells[0].merge(row.cells[1])
                row.cells[1].merge(row.cells[2])
                row.cells[2].text="Всего в "+str_(sem.number)+" семестре";
                row.cells[3].text=prepair_value_to_str(le_s+la_s+pa_s+sa_s+kr_pa_s+co_s)
                row.cells[4].text=prepair_value_to_str(le_s+la_s+pa_s)
                row.cells[5].text=prepair_value_to_str(le_s)
                row.cells[6].text=prepair_value_to_str(la_s)
                row.cells[7].text=prepair_value_to_str(pa_s)
                row.cells[8].text=prepair_value_to_str(sa_s)
                row.cells[10].text=prepair_value_to_str(sem.hours_summary.control);
                row.cells[9].text=prepair_value_to_str(sem.krpa);
                format_row(row,9)
                le_ss+=le_s
                la_ss+=la_s
                pa_ss+=pa_s
                sa_ss+=sa_s
                co_ss+=co_s
                kr_pa_ss+=kr_pa_s;
            row = temp_table.add_row()
            row.cells[0].merge(row.cells[1])
            row.cells[1].merge(row.cells[2])
            row.cells[2].text="Всего";
            row.cells[3].text=prepair_value_to_str(le_ss+la_ss+pa_ss+sa_ss+kr_pa_ss+co_ss)
            row.cells[4].text=prepair_value_to_str(le_ss+la_ss+pa_ss)
            row.cells[5].text=prepair_value_to_str(le_ss)
            row.cells[6].text=prepair_value_to_str(la_ss)
            row.cells[7].text=prepair_value_to_str(pa_ss)
            row.cells[8].text=prepair_value_to_str(sa_ss)
            row.cells[10].text=prepair_value_to_str(co_ss)
            row.cells[9].text=prepair_value_to_str(kr_pa_ss)
            format_row(row,9)
            for row in temp_table.rows:
                format_cell(row.cells[11],9)
            temp_table.style='Table Grid'
            
    for t in for_remove:
        t._element.getparent().remove(t._element)            
#def add_text_content(doc,pattern,text):

        
        
def get_comp_text(disc):
    UK=False;
    OPK=False;
    PK=False;
    UKs='';
    OPKs='';
    PKs='';
    text=""
    countUK=0
    countPK=0
    countOPK=0
    for comp in disc.competentions:
        if (re.match('^ОПК',comp.code)):
            OPK=True;
            countOPK+=1;
            if len(OPKs)>0:
                OPKs+=', '+comp.code
            else:
                OPKs+=comp.code
                
        if (re.match('^УК',comp.code)):
            UK=True;
            countUK+=1;
            if len(UKs)>0:
                UKs+=', '+comp.code
            else:
                UKs+=comp.code
                
        if (re.match('^ПК',comp.code)):
            PK=True;
            countPK+=1;
            if len(PKs)>0:
                PKs+=', '+comp.code
            else:
                PKs+=comp.code
    if countUK==1:
        text+='универсальной ('+UKs+')'
    elif countUK>1:
        text+='универсальных ('+UKs+')'
        
    if countOPK==1:
        if (len(text)>0):
            text+=', '
        text+='общепрофессиональной ('+OPKs+')'
    elif countOPK>1:
        if (len(text)>0):
            text+=', '
        text+='общепрофессиональных ('+OPKs+')'
        
    if countPK==1:
        if (len(text)>0):
            text+=', '
        text+='профессиональной ('+PKs+')'
    elif countPK>1:
        if (len(text)>0):
            text+=', '
        text+='профессиональных ('+PKs+')'
    #print(countUK)
    #print(countPK)
    #print(countOPK)
    if countUK+countPK+countOPK>1:
        text+=' компетенций'
    else:
        text+=' компетенции'
    return text;

def add_no_in_plan(doc):
    elements=[]
    for e in doc.element.iterchildren():
        for element in e.iterchildren():
            elements.append((element,e)) 
    indexes=range(0,len(elements)-1)
    for_return = dc.course_content();
    index=0
    indexes_to_insert =[]
    for i in indexes:
        if type(elements[i][0]) == docx.oxml.text.paragraph.CT_P:
            paragr = Paragraph(elements[i][0],elements[i][1])
            if re.match("4.4. П",paragr.text) is not None:
                if type(elements[i+1][0]) != docx.oxml.table.CT_Tbl:
                    indexes_to_insert.append(index+1)
            if re.match("4.3. Л",paragr.text) is not None:
                if type(elements[i+1][0]) != docx.oxml.table.CT_Tbl:
                    indexes_to_insert.append(index+1)
            index+=1
    for i in indexes_to_insert:
        doc.paragraphs[i].insert_paragraph_before('Учебным планом не предусмотрены');
    for i in indexes_to_insert:
        for run in doc.paragraphs[i].runs:
            run.font.size = docx.shared.Pt(14);
            run.font.name='Times New Roman';




def create_wp(disc,course_type):
    doc = docx.Document(settings.cfg.path_to_common)
    add_content_to_tables(disc,doc)
    add_no_in_plan(doc)
    add_text_content(doc,'{0}',disc.name)
    add_text_content(doc,'{1}',get_comp_text(disc))
    s=0
    ze=0
    for sem in disc.semesters:
        s+=sem.hours_summary.get_summary()+sem.krpa
        ze+=sem.zach_ed
    add_text_content(doc,'{3}',str_(s))
    add_text_content(doc,'{2}',str_(ze))
    
    
    curr=disc.current_questions
    cont=disc.control_questions
    if (len(curr)==0 and len(cont)!=0 ):
        curr=copy.copy(cont)
        random.shuffle(curr)
    if (len(cont)==0 and len(curr)!=0 ):
        cont=copy.copy(curr)
        random.shuffle(cont)
        
    if len(curr)>10:
        add_paragraphs_content(doc,'{4}',random.sample(curr,k=10))
    else:
        add_paragraphs_content(doc,'{4}',curr)
    if len(curr)>5:
        add_paragraphs_content(doc,'{5}',random.sample(curr,k=5))
    else:
        add_paragraphs_content(doc,'{5}',curr)
        
    add_paragraphs_content(doc,'{6}',curr)
    
    if (len(disc.control_questions)!=0):   
        add_paragraphs_content(doc,'{7}',cont)
    else:
        add_paragraphs_content(doc,'{7}',cont)
        
    add_paragraphs_content(doc,'{8}',disc.main_lit)
    add_paragraphs_content(doc,'{9}',disc.additional_lit)
    
    
    inn_text=""
    if (disc.lec_check()):
        inn_text+='лекции'
        remove_pattern(doc,'{11}')
    else:
        remove_paragraphs_by_pattern(doc,'{11}')
        
    if (disc.pract_check()):
        if (len(inn_text)>0):
            inn_text+=', практические занятия'
        else:
            inn_text+='практические занятия'
        remove_pattern(doc,'{12}')
    else:
        remove_paragraphs_by_pattern(doc,'{12}')
    if (disc.lab_check()):
        if (len(inn_text)>0):
            inn_text+=', лабораторные работы'
        else:
            inn_text+='лабораторные работы'    
        remove_pattern(doc,'{13}')
    else:
        remove_paragraphs_by_pattern(doc,'{13}')
    if (disc.sam_check()):
        if (len(inn_text)>0):
            inn_text+=', а также самостоятельную работу'
        else:
            inn_text+='самостоятельную работу'  
            
    add_text_content(doc,'{10}',inn_text)
    
    zn=[]
    um=[]
    vl=[]
    for comp in disc.competentions:
        zn.append(comp.znat+" ("+comp.code+")");
        um.append(comp.umet+" ("+comp.code+")");
        vl.append(comp.vladet+" ("+comp.code+")");
    
    add_paragraphs_content(doc,'{14}',zn,'-')
    add_paragraphs_content(doc,'{15}',um,'-')
    add_paragraphs_content(doc,'{16}',vl,'-')
    
    doc.tables[2].cell(0,0).text=disc.name
    for p in doc.tables[2].cell(0,0).paragraphs:
        for r in p.runs:
            r.font.name='Times New Roman';
            r.font.bold=True;
            r.font.size=docx.shared.Pt(14);
            
    doc.tables[3].cell(0,1).text=disc.autor
    for p in doc.tables[3].cell(0,1).paragraphs:
        for r in p.runs:
            r.font.name='Times New Roman';
            r.font.bold=True;
            r.font.size=docx.shared.Pt(14);
    add_text_content(doc,'{17}',parse_code(disc.code))
    
    doc.save(settings.cfg.path_to_result_dir+course_type+'2019_'+disc.name+'.docx');
    
    
for d in disc_list:
    print(settings.cfg.course+'  '+d.name)
    create_wp(d,settings.cfg.course)